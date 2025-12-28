"""Document extractors for different file types."""

from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import List, Optional
import io


@dataclass
class ExtractedContent:
    """Represents extracted content from a document."""

    text: str
    page_number: Optional[int] = None
    start_time_seconds: Optional[int] = None
    end_time_seconds: Optional[int] = None
    metadata: dict = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class BaseExtractor(ABC):
    """Base class for document extractors."""

    @abstractmethod
    def extract(self, file_content: bytes, filename: str) -> List[ExtractedContent]:
        """Extract text content from file.

        Args:
            file_content: Raw file bytes
            filename: Original filename

        Returns:
            List of ExtractedContent objects
        """
        pass

    @staticmethod
    def get_extractor(source_type: str) -> "BaseExtractor":
        """Factory method to get appropriate extractor.

        Args:
            source_type: Type of source (pdf, docx, txt, video, notes)

        Returns:
            Appropriate extractor instance
        """
        extractors = {
            "pdf": PDFExtractor,
            "docx": DocxExtractor,
            "txt": TextExtractor,
            "notes": TextExtractor,
            "video": VideoExtractor,
        }
        extractor_class = extractors.get(source_type.lower(), TextExtractor)
        return extractor_class()


class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using pdfplumber."""

    def extract(self, file_content: bytes, filename: str) -> List[ExtractedContent]:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required for PDF extraction. Install with: pip install pdfplumber"
            )

        contents = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    contents.append(
                        ExtractedContent(
                            text=text.strip(),
                            page_number=page_num,
                            metadata={"total_pages": len(pdf.pages)},
                        )
                    )

        return contents


class DocxExtractor(BaseExtractor):
    """Extract text from DOCX files using python-docx."""

    def extract(self, file_content: bytes, filename: str) -> List[ExtractedContent]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX extraction. Install with: pip install python-docx"
            )

        doc = Document(io.BytesIO(file_content))

        # Combine paragraphs with logical breaks
        paragraphs = []
        current_section = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if this looks like a heading (typically short and may be bold/styled)
                if para.style and "heading" in para.style.name.lower():
                    # Save current section and start new
                    if current_section:
                        paragraphs.append("\n".join(current_section))
                        current_section = []
                    current_section.append(f"## {text}")
                else:
                    current_section.append(text)

        # Don't forget last section
        if current_section:
            paragraphs.append("\n".join(current_section))

        # Return as single extracted content or split by sections
        if paragraphs:
            full_text = "\n\n".join(paragraphs)
            return [
                ExtractedContent(
                    text=full_text, metadata={"paragraph_count": len(doc.paragraphs)}
                )
            ]

        return []


class TextExtractor(BaseExtractor):
    """Extract text from plain text files."""

    def extract(self, file_content: bytes, filename: str) -> List[ExtractedContent]:
        # Try different encodings
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        text = None

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            text = file_content.decode("utf-8", errors="ignore")

        if text.strip():
            return [ExtractedContent(text=text.strip())]

        return []


class VideoExtractor(BaseExtractor):
    """Extract text from video/audio files using Whisper transcription.

    Note: This is a placeholder. Full implementation requires whisper setup.
    For production, consider Azure Speech SDK for better scalability.
    """

    def extract(self, file_content: bytes, filename: str) -> List[ExtractedContent]:
        # Try to use whisper if available
        try:
            import whisper
            import tempfile
            import os

            # Write to temp file (whisper needs file path)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                model = whisper.load_model(
                    "base"
                )  # Use "small" or "medium" for better accuracy
                result = model.transcribe(tmp_path)

                contents = []
                # Extract segments with timestamps
                for segment in result.get("segments", []):
                    contents.append(
                        ExtractedContent(
                            text=segment["text"].strip(),
                            start_time_seconds=int(segment["start"]),
                            end_time_seconds=int(segment["end"]),
                            metadata={"language": result.get("language", "en")},
                        )
                    )

                return (
                    contents
                    if contents
                    else [ExtractedContent(text=result.get("text", ""))]
                )
            finally:
                os.unlink(tmp_path)

        except ImportError:
            # Whisper not installed - return placeholder
            return [
                ExtractedContent(
                    text=f"[Video transcription pending - whisper not installed. File: {filename}]",
                    metadata={"transcription_status": "pending"},
                )
            ]
        except Exception as e:
            return [
                ExtractedContent(
                    text=f"[Video transcription failed: {str(e)}. File: {filename}]",
                    metadata={"transcription_status": "failed", "error": str(e)},
                )
            ]
